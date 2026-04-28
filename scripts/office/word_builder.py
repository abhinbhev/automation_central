"""Generate Word documents from a structured JSON spec using python-docx.

Usage:
    python word_builder.py --spec doc_spec.json --output output.docx
    python word_builder.py --spec doc_spec.json --template sop --output sop.docx

The JSON spec format:
{
  "type": "sop",
  "title": "Document Title",
  "metadata": {
    "owner": "Team Name",
    "date": "2025-05-23",
    "version": "1.0"
  },
  "sections": [
    {
      "heading": "Section Heading",
      "level": 1,
      "content": "Section body text.",
      "table": {
        "headers": ["Column A", "Column B"],
        "rows": [["value1", "value2"]]
      }
    }
  ]
}

Templates in templates/word/ are used as base .docx if available.
"""

import json
from datetime import date
from pathlib import Path

import typer
from docx import Document
from rich.console import Console

app = typer.Typer()
console = Console()

TEMPLATES_DIR = Path(__file__).resolve().parents[2] / "templates" / "word"
HEADING_STYLE_MAP = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}


def get_template_path(template_name: str) -> Path | None:
    candidate = TEMPLATES_DIR / f"{template_name}.docx"
    if candidate.exists():
        return candidate
    console.print(f"[yellow]Template '{template_name}.docx' not found — using blank document.[/yellow]")
    return None


def add_metadata_table(doc: Document, metadata: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = "Value"
    for key, value in metadata.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key.replace("_", " ").title()
        row_cells[1].text = str(value)
    doc.add_paragraph()


def add_section(doc: Document, section: dict) -> None:
    heading = section.get("heading", "")
    level = section.get("level", 1)
    content = section.get("content", "")
    table_data = section.get("table")
    items = section.get("items", [])  # bulleted list

    style = HEADING_STYLE_MAP.get(level, "Heading 1")
    doc.add_heading(heading, level=level)

    if content:
        doc.add_paragraph(content)

    if items:
        for item in items:
            para = doc.add_paragraph(style="List Bullet")
            para.text = item

    if table_data:
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        if headers:
            tbl = doc.add_table(rows=1, cols=len(headers))
            tbl.style = "Table Grid"
            hdr_cells = tbl.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs[0] if hdr_cells[i].paragraphs[0].runs else hdr_cells[i].paragraphs[0].add_run(h)
                run.bold = True
            for row in rows:
                row_cells = tbl.add_row().cells
                for i, val in enumerate(row):
                    if i < len(row_cells):
                        row_cells[i].text = str(val)
            doc.add_paragraph()


def build_document(spec: dict, template_name: str, output_path: Path) -> None:
    template_path = get_template_path(template_name)
    doc = Document(str(template_path)) if template_path else Document()

    title = spec.get("title", "Untitled Document")
    metadata = spec.get("metadata", {})
    sections = spec.get("sections", [])

    doc.add_heading(title, 0)

    if metadata:
        if "date" not in metadata:
            metadata["date"] = str(date.today())
        add_metadata_table(doc, metadata)

    for section in sections:
        add_section(doc, section)

    doc.save(str(output_path))


@app.command()
def main(
    spec: Path = typer.Option(..., help="Path to JSON document spec"),
    output: Path = typer.Option(..., help="Output .docx file path"),
    template: str = typer.Option("default", help="Template name from templates/word/"),
) -> None:
    """Generate a Word document from a structured JSON spec."""
    if not spec.exists():
        console.print(f"[red]Spec file not found:[/red] {spec}")
        raise typer.Exit(1)

    doc_spec = json.loads(spec.read_text())
    if not isinstance(doc_spec, dict):
        console.print("[red]Spec file must be a JSON object.[/red]")
        raise typer.Exit(1)

    console.print(f"Building document '{doc_spec.get('title', 'Untitled')}'...")
    build_document(doc_spec, template, output)
    console.print(f"[green]✓ Document saved to {output}[/green]")


if __name__ == "__main__":
    app()
