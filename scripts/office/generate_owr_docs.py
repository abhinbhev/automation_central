"""Generate OWR + HLR Word documents for the 3 ADO features.

Reads docs/156057-*.md, docs/156061-*.md, docs/156063-*.md and produces
styled .docx files in docs/ following the OWR+HLR document template pattern.

Usage:
    python scripts/office/generate_owr_docs.py
    python scripts/office/generate_owr_docs.py --out-dir docs/word
"""

from __future__ import annotations

from pathlib import Path

import typer
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from rich.console import Console

app = typer.Typer()
console = Console()

REPO_ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = REPO_ROOT / "docs"

# ABI brand colours
YELLOW = RGBColor(0xF5, 0xC5, 0x18)   # ABI gold/yellow
DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
MID    = RGBColor(0x4A, 0x4A, 0x6A)   # mid grey-blue
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
HEADER_BG  = RGBColor(0x1A, 0x1A, 0x2E)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, rgb: RGBColor) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_run_color(run, rgb: RGBColor) -> None:
    run.font.color.rgb = rgb


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "F5C518")
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Document builders
# ---------------------------------------------------------------------------

def _add_cover_header(doc: Document, feature_id: str, title: str, status: str, owner: str, version: str, updated: str) -> None:
    """Dark banner header with title and metadata strip."""
    # Title paragraph with dark background simulation via shading on a table
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = "Table Grid"

    # Row 0 — title banner
    title_cell = tbl.rows[0].cells[0]
    _set_cell_bg(title_cell, HEADER_BG)
    title_para = title_cell.paragraphs[0]
    title_para.clear()
    run = title_para.add_run(f"OWR + HLR")
    run.bold = True
    run.font.size = Pt(10)
    _set_run_color(run, YELLOW)
    title_para.add_run("\n")
    run2 = title_para.add_run(title)
    run2.bold = True
    run2.font.size = Pt(16)
    _set_run_color(run2, WHITE)
    title_para.paragraph_format.space_before = Pt(10)
    title_para.paragraph_format.space_after = Pt(10)
    title_para.paragraph_format.left_indent = Inches(0.2)

    # Row 1 — metadata strip
    meta_cell = tbl.rows[1].cells[0]
    _set_cell_bg(meta_cell, LIGHT_GREY)
    meta_para = meta_cell.paragraphs[0]
    meta_para.clear()
    meta_text = f"ADO Feature: {feature_id}   |   Status: {status}   |   Owner: {owner}   |   Version: {version}   |   Last updated: {updated}"
    run3 = meta_para.add_run(meta_text)
    run3.font.size = Pt(8)
    run3.font.color.rgb = MID
    meta_para.paragraph_format.space_before = Pt(4)
    meta_para.paragraph_format.space_after = Pt(4)
    meta_para.paragraph_format.left_indent = Inches(0.2)

    doc.add_paragraph()


def _add_section1_table(doc: Document, fields: dict[str, str]) -> None:
    """Section 1 — Product Brief info table."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(2.0)
    tbl.columns[1].width = Inches(4.5)

    hdr = tbl.rows[0].cells
    for cell in hdr:
        _set_cell_bg(cell, HEADER_BG)
    hdr[0].paragraphs[0].clear()
    r = hdr[0].paragraphs[0].add_run("Field")
    r.bold = True; r.font.size = Pt(9); _set_run_color(r, WHITE)
    hdr[1].paragraphs[0].clear()
    r = hdr[1].paragraphs[0].add_run("Value")
    r.bold = True; r.font.size = Pt(9); _set_run_color(r, WHITE)

    for key, val in fields.items():
        row = tbl.add_row().cells
        row[0].paragraphs[0].clear()
        rk = row[0].paragraphs[0].add_run(key)
        rk.bold = True; rk.font.size = Pt(9)
        row[1].paragraphs[0].clear()
        rv = row[1].paragraphs[0].add_run(val)
        rv.font.size = Pt(9)

    doc.add_paragraph()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_heading(text, level=level)
    # Tint H1 with brand yellow accent
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = DARK
    if level == 2:
        for run in p.runs:
            run.font.color.rgb = MID


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"

    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_bg(hdr_cells[i], HEADER_BG)
        hdr_cells[i].paragraphs[0].clear()
        r = hdr_cells[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        _set_run_color(r, WHITE)

    for row_data in rows:
        row_cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].paragraphs[0].clear()
                r = row_cells[i].paragraphs[0].add_run(str(val))
                r.font.size = Pt(9)

    doc.add_paragraph()


def _add_bullet_list(doc: Document, items: list[str], bold_prefix: str | None = None) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if bold_prefix and item.startswith(bold_prefix):
            parts = item.split(":", 1)
            if len(parts) == 2:
                run = p.add_run(parts[0] + ":")
                run.bold = True
                run.font.size = Pt(9)
                run2 = p.add_run(parts[1])
                run2.font.size = Pt(9)
                continue
        run = p.add_run(item)
        run.font.size = Pt(9)


def _add_checklist_section(doc: Document, phase: str, items: list[str], required_output: str) -> None:
    _add_heading(doc, phase, level=3)
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"☐  {item}")
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    run = p.add_run("Required output: ")
    run.bold = True
    run.font.size = Pt(9)
    run2 = p.add_run(required_output)
    run2.font.size = Pt(9)
    run2.italic = True
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Document specs
# ---------------------------------------------------------------------------

def _build_156057(doc: Document) -> None:
    """ADO 156057 — Product Analytics (IC Component Tracking)."""
    _add_cover_header(
        doc,
        feature_id="156057 — Product Analytics",
        title="Product Analytics (IC Component Tracking)",
        status="DRAFT",
        owner="POS IC X OneWay",
        version="0.1",
        updated="2026-04-30",
    )

    # Section 1
    _add_heading(doc, "Section 1 — Product Brief")
    _add_section1_table(doc, {
        "ADO Feature": "156057 — Product Analytics",
        "Area Path": "Generative AI Products \\ POS IC X OneWay",
        "Target Date": "TBD",
        "Document Status": "DRAFT",
        "Team Members": "TBD",
        "Design File": "TBD",
        "ADO Backlog": "https://dev.azure.com/ab-inbev-analytics",
    })

    _add_horizontal_rule(doc)

    # Section 2
    _add_heading(doc, "Section 2 — Feature Body")

    _add_heading(doc, "Objective", level=2)
    _add_bullet_list(doc, [
        "Implement comprehensive Google Analytics 4 (GA4) tracking for all user interactions with the IC (Insights Copilot / 'Ask OneWay AI') component in OneWay.",
        "Enable product managers to measure adoption, understand usage patterns, identify popular features, and improve the product based on data-driven insights.",
        "Track 21 distinct user interaction events across the full IC component lifecycle — from first impression through feedback and conversation management.",
    ])
    doc.add_paragraph()

    _add_heading(doc, "Success Metrics", level=2)
    _add_table(doc,
        headers=["Goal", "Metric"],
        rows=[
            ["Adoption", "ic_component_visible event fires once per new user (lifetime deduplication)"],
            ["Engagement", "Volume of ic_component_open / ic_component_close events over time"],
            ["Engagement", "Distribution of chat_submitted by mode (FAQ, Smart Flow, enter_key, button_click)"],
            ["Engagement", "Volume of chat_submitted with a conversation_id (follow-up rate)"],
            ["Outcome", "chat_response_received breakdown by status (SUCCESS / FAILURE / ABANDONED)"],
            ["Outcome", "track_feedback volume and feedback_sentiment (true / false) ratio"],
            ["Outcome", "track_feedback_submit options_count distribution (depth of feedback)"],
            ["Business Impact", "Data-driven product decisions enabled by full funnel GA4 event coverage"],
        ]
    )

    _add_heading(doc, "Assumptions", level=2)
    _add_heading(doc, "Features in scope", level=3)
    _add_bullet_list(doc, [
        "All 21 GA4 event scenarios defined in the acceptance criteria",
        "IC component within the OneWay application only",
        "currentProjectKey hardcoded to gai_copilot_marketing_brand_guidance_ghq for all events",
    ])
    _add_heading(doc, "Technical approach", level=3)
    _add_bullet_list(doc, [
        "Use Google Analytics 4 (GA4) event tracking via gtag.js or analytics.js",
        "conversation_id and chat_id extracted from Clara API responses",
        "ic_component_visible fired once per user — deduplicated via localStorage or cookie",
        "conversation_id propagated through all follow-up messages in a session",
    ])
    _add_heading(doc, "Data scope", level=3)
    _add_table(doc,
        headers=["Dimension", "Value"],
        rows=[
            ["Brands", "All brands using the IC component in OneWay"],
            ["Regions", "All regions where OneWay is deployed"],
            ["Channels", "Web (OneWay application)"],
            ["Periods", "From release date onwards"],
            ["Granularity", "Per user interaction event"],
        ]
    )

    _add_heading(doc, "Milestones", level=2)
    _add_table(doc,
        headers=["Milestone", "Date"],
        rows=[
            ["ADO Feature created", "—"],
            ["Development start", "TBD"],
            ["QA / validation", "TBD"],
            ["Release", "TBD"],
        ]
    )

    _add_heading(doc, "Requirements", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Event catalogue — all events must fire with ")
    r.font.size = Pt(9)
    r2 = p.add_run("currentProjectKey = gai_copilot_marketing_brand_guidance_ghq")
    r2.bold = True; r2.font.size = Pt(9)
    doc.add_paragraph()

    _add_table(doc,
        headers=["Event Name", "Trigger", "Key Parameters"],
        rows=[
            ["ic_component_visible", "IC component first renders on screen", "currentProjectKey — fires once per user lifetime"],
            ["ic_component_open", "User opens the 'Ask OneWay AI' panel", "currentProjectKey — fires every open"],
            ["ic_component_close", "User closes the 'Ask OneWay AI' panel", "currentProjectKey — fires every close"],
            ["chat_submitted", "User submits via FAQ (landing page)", "currentProjectKey, mode=FAQ, conversation_id (follow-up only)"],
            ["chat_submitted", "User selects Smart Flow from landing page", "currentProjectKey, mode=is_smartflow_landing_page, smartflow_function, conversation_id (follow-up only)"],
            ["chat_submitted", "User selects Smart Flow from chat bar", "currentProjectKey, mode=is_smartflow_chatbar, smartflow_function, conversation_id (follow-up only)"],
            ["chat_submitted", "User submits via Enter key", "currentProjectKey, mode=enter_key, conversation_id (follow-up only)"],
            ["chat_submitted", "User clicks Send button", "currentProjectKey, mode=button_click, conversation_id (follow-up only)"],
            ["chat_response_received", "Clara returns a successful response (task_status=success)", "currentProjectKey, conversation_id, chat_id, status=SUCCESS"],
            ["chat_response_received", "Clara fails to respond (task_status=failure)", "currentProjectKey, conversation_id, chat_id, status=FAILURE"],
            ["chat_response_received", "Response abandoned (task_status=abandoned)", "currentProjectKey, conversation_id, chat_id, status=ABANDONED"],
            ["track_new_chat", "User clicks 'New Chat' icon", "currentProjectKey"],
            ["track_feedback", "User clicks 👍", "currentProjectKey, feedback_sentiment=true, source=thumbs_icon"],
            ["track_feedback", "User clicks 👎", "currentProjectKey, feedback_sentiment=false, source=thumbs_icon"],
            ["track_feedback_submit", "User submits positive feedback options", "currentProjectKey, sentiment=true, options (comma-separated), options_count (0–6)"],
            ["track_feedback_submit", "User submits negative feedback options", "currentProjectKey, sentiment=false, options (comma-separated), options_count (0–4)"],
            ["rename_button_click", "User clicks Rename button", "currentProjectKey, conversation_id"],
            ["rename_text_change", "User confirms new conversation name", "currentProjectKey, conversation_id"],
            ["delete_confirm_button_click", "User clicks Confirm on delete modal", "currentProjectKey, conversation_id"],
            ["delete_cancel_button_click", "User clicks Cancel on delete modal", "currentProjectKey, conversation_id"],
            ["copy_code_button_click", "User clicks Copy icon on a code block", "currentProjectKey, conversation_id"],
        ]
    )

    _add_heading(doc, "Business rules", level=3)
    _add_bullet_list(doc, [
        "ic_component_visible must fire only once per user across all sessions — use localStorage or a persistent cookie to gate the event.",
        "conversation_id must be propagated from the first message through all follow-up messages in the same conversation.",
        "options_count for positive feedback ranges 0–6 (options: Accuracy, Relevance, Visuals, Completion, Language, Others).",
        "options_count for negative feedback ranges 0–4 (options: Accuracy, Relevance, Completion, Others).",
        "mode for chat_submitted must be one of exactly: FAQ, is_smartflow_landing_page, is_smartflow_chatbar, enter_key, button_click.",
        "status for chat_response_received must be one of exactly: SUCCESS, FAILURE, ABANDONED — derived from task_status in the Clara API response.",
    ])

    _add_heading(doc, "Out of Scope", level=2)
    _add_heading(doc, "Short-term", level=3)
    _add_bullet_list(doc, [
        "Custom GA4 dashboards or report configuration",
        "Tracking interactions outside the IC component (e.g., standard OneWay navigation events)",
        "Server-side event tracking (all tracking is client-side via gtag.js / analytics.js)",
    ])
    _add_heading(doc, "Long-term", level=3)
    _add_bullet_list(doc, [
        "Cross-project currentProjectKey dynamic configuration",
        "Integration with non-GA4 analytics platforms",
    ])

    _add_heading(doc, "Design", level=2)
    _add_table(doc,
        headers=["Artefact", "Link"],
        rows=[["Design file", "TBD"], ["Event schema reference", "TBD"]]
    )

    _add_heading(doc, "Open Questions", level=2)
    _add_table(doc,
        headers=["Question", "Answer", "Date Answered"],
        rows=[
            ["What is the exact GA4 property / measurement ID to use?", "—", "—"],
            ["Should ic_component_visible use localStorage or a cookie — and what is the TTL?", "—", "—"],
            ["Who owns the GA4 property and will validate event receipt?", "—", "—"],
            ["Is conversation_id always available in the API response from the first message, or only on follow-up?", "—", "—"],
        ]
    )

    _add_heading(doc, "Change & Request Log", level=2)
    _add_table(doc,
        headers=["Request / Update", "Requestor / Personnel", "Date"],
        rows=[["Initial draft created from ADO 156057", "—", "2026-04-30"]]
    )

    _add_heading(doc, "Reference Links", level=2)
    _add_bullet_list(doc, [
        "ADO Feature 156057: https://dev.azure.com/ab-inbev-analytics",
        "GA4 Event tracking: https://developers.google.com/analytics/devguides/collection/ga4/events",
    ])

    _add_horizontal_rule(doc)

    # Section 3
    _add_heading(doc, "Section 3 — Discovery → Delivery Checklist")

    _add_checklist_section(doc, "Phase 1 — Discovery Framing & Alignment", [
        "Problem statement defined: measure IC component adoption and usage via GA4",
        "Business intent confirmed: data-driven product improvements for IC",
        "Scope agreed: 21 events, client-side, GA4, OneWay only",
        "Ownership assigned (PM, engineering lead, analytics)",
        "ADO feature linked and prioritised",
    ], "Signed-off problem statement + scope boundary")

    _add_checklist_section(doc, "Phase 2 — Business Rules & Data Contract", [
        "All 21 event names and parameter schemas locked (see Requirements table)",
        "ic_component_visible deduplication mechanism confirmed (localStorage vs cookie, TTL)",
        "conversation_id propagation contract confirmed with Clara API team",
        "mode enum values locked for chat_submitted",
        "status enum values locked for chat_response_received",
        "options and options_count ranges confirmed for feedback events",
        "GA4 property / measurement ID confirmed",
        "currentProjectKey value confirmed as gai_copilot_marketing_brand_guidance_ghq",
    ], "Event schema document / data contract signed off")

    _add_checklist_section(doc, "Phase 3 — Discovery Validation (Users)", [
        "Key user journeys validated against the 21 scenarios",
        "Edge cases confirmed: abandoned responses, empty feedback submissions, rename/delete cancellations",
    ], "Validated scenario list with no gaps")

    _add_checklist_section(doc, "Phase 4 — HLR (High-Level Requirements)", [
        "MVP scope confirmed: all 21 events in initial release",
        "Non-goals documented (no GA4 dashboards, no server-side tracking)",
        "Dependencies identified: Clara API (conversation_id, chat_id, task_status), GA4 property access",
        "gtag.js / analytics.js library integration approach agreed",
    ], "This document (HLR) approved")

    _add_checklist_section(doc, "Phase 5 — Design Handoff & Prototype Validation", [
        "Event firing points mapped to UI components (buttons, icons, panel open/close)",
        "No design changes required — tracking is additive/non-visual",
    ], "Component-to-event mapping reviewed by engineering")

    _add_checklist_section(doc, "Phase 6 — Delivery Prep (User Stories & Walkthrough)", [
        "User stories created in ADO for each event group (visibility, open/close, chat submission, response, feedback, conversation management, copy)",
        "Acceptance criteria include: event name, parameters, trigger conditions, deduplication rules",
        "Empty / error state handling included (e.g., FAILURE, ABANDONED statuses)",
        "GA4 debug view validation approach agreed",
    ], "ADO stories with full AC, ready for sprint")

    _add_checklist_section(doc, "Phase 7 — Release & Communication", [
        "Release note drafted covering new GA4 event instrumentation",
        "Known limitations documented (client-side only, single currentProjectKey)",
        "Stakeholders notified (PM, analytics team, data consumers)",
        "GA4 DebugView or Tag Assistant validation completed",
    ], "Release note + stakeholder comms sent")

    _add_checklist_section(doc, "Phase 8 — Post-Release Validation & Learning", [
        "Confirm all 21 events are firing in production GA4 property",
        "Confirm ic_component_visible deduplication is working (no repeat fires per user)",
        "Confirm conversation_id propagation is consistent across multi-turn conversations",
        "Adoption and engagement KPIs baselined",
        "v2 improvements identified from initial data (if any)",
    ], "GA4 event validation report + adoption baseline")


def _build_156061(doc: Document) -> None:
    """ADO 156061 — Share (Chat, Conversation, Conversation via History)."""
    _add_cover_header(
        doc,
        feature_id="156061 — Share (Chat, Conversation, Conversation via History)",
        title="Share (Chat, Conversation, Conversation via History)",
        status="DRAFT",
        owner="POS IC X OneWay",
        version="0.1",
        updated="2026-04-30",
    )

    _add_heading(doc, "Section 1 — Product Brief")
    _add_section1_table(doc, {
        "ADO Feature": "156061 — Share (Chat, Conversation, Conversation via History)",
        "Area Path": "Generative AI Products \\ POS IC X OneWay",
        "Target Date": "TBD",
        "Document Status": "DRAFT",
        "Team Members": "TBD",
        "Design File": "TBD",
        "ADO Backlog": "https://dev.azure.com/ab-inbev-analytics",
    })

    _add_horizontal_rule(doc)

    _add_heading(doc, "Section 2 — Feature Body")

    _add_heading(doc, "Objective", level=2)
    _add_bullet_list(doc, [
        "Allow Clara users to share individual responses and full conversations with colleagues via a shareable link.",
        "Enable recipients to view shared content in read-only mode and optionally fork the conversation as a new unique instance in their own history.",
        "Expose sharing from three entry points: the response action bar, the active conversation view, and the conversation history panel.",
        "Ensure shared content respects recipient RBAC permissions — data visibility is scoped to what the recipient is authorised to see.",
    ])
    doc.add_paragraph()

    _add_heading(doc, "Success Metrics", level=2)
    _add_table(doc,
        headers=["Goal", "Metric"],
        rows=[
            ["Adoption", "Number of share link generation events per week"],
            ["Adoption", "% of active Clara users who use the share feature at least once per month"],
            ["Engagement", "Number of shared links opened by recipients"],
            ["Engagement", "% of recipients who fork a shared conversation into their own history"],
            ["Outcome", "Reduction in duplicate queries (colleagues querying the same question independently)"],
            ["Outcome", "Zero access control violations — recipients never see data beyond their RBAC permissions"],
            ["Business Impact", "Faster team-level decision making using shared Clara insights"],
        ]
    )

    _add_heading(doc, "Assumptions", level=2)
    _add_heading(doc, "Features in scope", level=3)
    _add_bullet_list(doc, [
        "Share a single response via link (from response action bar)",
        "Share an entire conversation via link (from active conversation view header)",
        "Share a conversation from the history panel via '...' context menu",
        "Recipients can view shared content in read-only mode",
        "Recipients with OneWay access can fork any shared conversation as a unique instance in their own conversation history",
        "Recipients cannot edit or delete the sharer's original conversation instance",
        "RBAC permissions apply to data visibility for recipients",
    ])
    _add_heading(doc, "Technical approach", level=3)
    _add_bullet_list(doc, [
        "Shareable links generated server-side; links encode the conversation or response identifier",
        "Recipient access enforced at load time — if RBAC prevents data access, the recipient sees restricted content or an access error",
        "Forking creates a new independent conversation in the recipient's history — does not modify the original",
        "The share link icon (🔗) and 'Share Conversation' option are additive to existing UI components; no existing functionality is removed",
    ])
    _add_heading(doc, "Data scope", level=3)
    _add_table(doc,
        headers=["Dimension", "Value"],
        rows=[
            ["Brands", "All brands accessible via the user's OneWay RBAC permissions"],
            ["Regions", "All regions accessible via the user's OneWay RBAC permissions"],
            ["Channels", "Web (OneWay application)"],
            ["Periods", "From release date onwards"],
            ["Granularity", "Per response / per conversation"],
        ]
    )

    _add_heading(doc, "Milestones", level=2)
    _add_table(doc,
        headers=["Milestone", "Date"],
        rows=[
            ["ADO Feature created", "—"],
            ["Development start", "TBD"],
            ["QA / validation", "TBD"],
            ["Release", "TBD"],
        ]
    )

    _add_heading(doc, "Requirements", level=2)
    _add_heading(doc, "Scenario 1 — Share a single response via link", level=3)
    _add_bullet_list(doc, [
        "A share link icon (🔗) must be visible in the response action bar alongside 👍, 👎, and 📋 (copy) icons.",
        "Clicking 🔗 generates a shareable link to the specific response.",
        "The user can copy the link or share it via Teams/email.",
        "Recipients with OneWay access can view the response in read-only mode.",
        "Recipients with OneWay access can continue the conversation as a unique instance in their own history.",
    ])
    _add_heading(doc, "Scenario 2 — Share entire conversation via link", level=3)
    _add_bullet_list(doc, [
        "A 'Share Conversation' option must be visible at the top of the active chat window.",
        "Clicking it generates a shareable read-only link for the full conversation session.",
        "The user can copy the link or send it via email/chat.",
        "Recipients with OneWay access can view the full conversation in read-only mode.",
        "Recipients with OneWay access can continue the conversation as a unique instance in their own history.",
    ])
    _add_heading(doc, "Scenario 3 — Share from conversation history panel", level=3)
    _add_bullet_list(doc, [
        "In the conversation history panel, each conversation must have a '...' context menu.",
        "The '...' menu must include a 'Share Conversation' option.",
        "Clicking 'Share Conversation' generates a shareable link for that specific conversation session.",
        "The user can copy the link or share it via email/chat.",
        "Recipients with OneWay access can continue the conversation as a unique instance in their own history.",
    ])
    _add_heading(doc, "Scenario 4 — Shared conversation access control", level=3)
    _add_bullet_list(doc, [
        "Recipients can view the shared conversation or response in read-only mode only.",
        "Recipients can fork and continue the conversation as a unique, independent instance in their own history.",
        "Recipients cannot edit or delete the sharer's instance or copy of the conversation.",
        "Data visibility for recipients must respect their RBAC permissions — recipients only see data they are authorised to access.",
    ])

    _add_heading(doc, "Out of Scope", level=2)
    _add_heading(doc, "Short-term", level=3)
    _add_bullet_list(doc, [
        "Public / unauthenticated share links (all recipients must have OneWay access)",
        "Share link expiry or revocation controls",
        "Notification to the original sharer when a recipient opens or forks a shared link",
        "Sharing to external users outside ABI / outside OneWay",
    ])
    _add_heading(doc, "Long-term", level=3)
    _add_bullet_list(doc, [
        "Collaborative real-time editing of a shared conversation",
        "Organisation-wide or team-wide shared conversation libraries",
    ])

    _add_heading(doc, "Design", level=2)
    _add_table(doc,
        headers=["Artefact", "Link"],
        rows=[
            ["Design file", "TBD"],
            ["Response action bar mockup (with 🔗 icon)", "TBD"],
            ["'Share Conversation' header option mockup", "TBD"],
            ["History panel '...' menu mockup", "TBD"],
            ["Shared conversation read-only view mockup", "TBD"],
        ]
    )

    _add_heading(doc, "Open Questions", level=2)
    _add_table(doc,
        headers=["Question", "Answer", "Date Answered"],
        rows=[
            ["What happens when a recipient opens a shared link but lacks RBAC access to the underlying data — partial view or full access error?", "—", "—"],
            ["Is the 'fork to my history' action automatic on link open, or does the recipient explicitly click 'Continue conversation'?", "—", "—"],
            ["Are share links permanent or do they expire? If they expire, what is the TTL?", "—", "—"],
            ["Can the original sharer revoke a share link after sending?", "—", "—"],
            ["Is a shared single response link a deep link into a conversation, or a standalone response page?", "—", "—"],
            ["Should GA4 tracking events be added for share link generation and link open? (See Feature 156057)", "—", "—"],
            ["What is the exact UI label for the share option in the conversation header — 'Share Conversation' or something else?", "—", "—"],
        ]
    )

    _add_heading(doc, "Change & Request Log", level=2)
    _add_table(doc,
        headers=["Request / Update", "Requestor / Personnel", "Date"],
        rows=[["Initial draft created from ADO 156061", "—", "2026-04-30"]]
    )

    _add_heading(doc, "Reference Links", level=2)
    _add_bullet_list(doc, [
        "ADO Feature 156061: https://dev.azure.com/ab-inbev-analytics",
        "Related: ADO 156057 (Product Analytics — share events may need GA4 tracking)",
        "Related: ADO 156063 (Domain Differentiation — Brand context applies to shared conversations)",
    ])

    _add_horizontal_rule(doc)

    _add_heading(doc, "Section 3 — Discovery → Delivery Checklist")

    _add_checklist_section(doc, "Phase 1 — Discovery Framing & Alignment", [
        "Problem statement defined: users cannot share Clara insights with colleagues without re-querying",
        "Business intent confirmed: enable team collaboration on data-driven decisions via shareable links",
        "Scope agreed: 4 scenarios (single response share, full conversation share, history panel share, access control)",
        "Ownership assigned (PM, UX, engineering, backend/API)",
        "ADO feature linked and prioritised",
    ], "Signed-off problem statement + scope boundary")

    _add_checklist_section(doc, "Phase 2 — Business Rules & Data Contract", [
        "Share link structure defined (response vs. conversation — same or different endpoint/format)",
        "RBAC enforcement mechanism confirmed: how recipient data scoping is applied at link open time",
        "Fork behaviour confirmed: what exactly is copied to recipient history (full conversation, from a specific message?)",
        "Read-only mode behaviour defined: what actions are disabled for recipients",
        "Link permanence / expiry policy confirmed",
        "Revocation capability confirmed (in or out of scope)",
        "GA4 tracking for share events — confirm with Feature 156057 team",
    ], "Share link technical spec + RBAC contract")

    _add_checklist_section(doc, "Phase 3 — Discovery Validation (Users)", [
        "Users confirm the 3 share entry points cover their sharing workflows",
        "Read-only + fork model validated as sufficient for recipient use cases",
        "RBAC restriction behaviour validated — recipients accept that restricted data will not be visible",
    ], "User validation notes")

    _add_checklist_section(doc, "Phase 4 — HLR (High-Level Requirements)", [
        "MVP scope confirmed: 4 scenarios, all 3 entry points, RBAC enforcement, fork to history",
        "Non-goals documented: no public links, no link expiry, no revocation, no external sharing",
        "Dependencies identified: share link generation API, RBAC enforcement at read time, conversation history service, link routing",
        "Fork mechanism dependency on conversation history service confirmed",
    ], "This document (HLR) approved")

    _add_checklist_section(doc, "Phase 5 — Design Handoff & Prototype Validation", [
        "Response action bar with 🔗 icon reviewed",
        "'Share Conversation' header option reviewed",
        "History panel '...' menu with share option reviewed",
        "Shared conversation read-only view reviewed (which actions are hidden/disabled)",
        "RBAC restriction state reviewed (partial data view vs. access denied message)",
    ], "Design sign-off from PO / UX lead")

    _add_checklist_section(doc, "Phase 6 — Delivery Prep (User Stories & Walkthrough)", [
        "User stories created in ADO for: single response share, full conversation share, history panel share, access control / RBAC enforcement, fork to history",
        "Acceptance criteria include: link generation, read-only mode restrictions, fork behaviour, RBAC data scoping",
        "Error states documented: invalid/expired link, insufficient RBAC access, conversation no longer available",
        "Empty states documented: conversation with no messages, single response with no conversation context",
    ], "ADO stories with full AC, ready for sprint")

    _add_checklist_section(doc, "Phase 7 — Release & Communication", [
        "Release note drafted covering sharing feature (3 entry points, read-only + fork model)",
        "Known limitations documented (OneWay access required, no link expiry/revocation, no external sharing)",
        "Stakeholders notified (PM, OneWay users, teams lead)",
        "Enablement material created if needed (short guide or tooltip copy for share UX)",
    ], "Release note + stakeholder comms sent")

    _add_checklist_section(doc, "Phase 8 — Post-Release Validation & Learning", [
        "Confirm share link generation works from all 3 entry points in production",
        "Confirm read-only mode correctly disables edit/delete for recipients",
        "Confirm fork to history creates independent conversation instances",
        "Confirm RBAC data scoping is enforced correctly for recipients with different permission levels",
        "Monitor share link generation and open rates (KPIs from Success Metrics)",
        "v2 improvements identified (e.g., link expiry, revocation, notification to sharer)",
    ], "Post-release validation checklist signed off")


def _build_156063(doc: Document) -> None:
    """ADO 156063 — Domain Differentiation (Clara Brand Guidance)."""
    _add_cover_header(
        doc,
        feature_id="156063 — Domain Differentiation",
        title="Domain Differentiation (Clara Brand Guidance)",
        status="DRAFT",
        owner="POS IC X OneWay",
        version="0.1",
        updated="2026-04-30",
    )

    _add_heading(doc, "Section 1 — Product Brief")
    _add_section1_table(doc, {
        "ADO Feature": "156063 — Domain Differentiation",
        "Area Path": "Generative AI Products \\ POS IC X OneWay",
        "Target Date": "TBD",
        "Document Status": "DRAFT",
        "Team Members": "TBD",
        "Design File": "TBD",
        "ADO Backlog": "https://dev.azure.com/ab-inbev-analytics",
    })

    _add_horizontal_rule(doc)

    _add_heading(doc, "Section 2 — Feature Body")

    _add_heading(doc, "Objective", level=2)
    _add_bullet_list(doc, [
        "Position Clara clearly as a Brand Guidance-focused assistant within the OneWay UI and its underlying behaviour.",
        "Ensure users understand Clara's scope on first interaction so they ask relevant Brand-related questions and receive context-aware responses.",
        "Prevent misleading or irrelevant responses to out-of-scope queries by having Clara handle them gracefully.",
        "Enforce Brand context through deep links, redirecting users to the correct page before loading a shared conversation.",
    ])
    doc.add_paragraph()

    _add_heading(doc, "Success Metrics", level=2)
    _add_table(doc,
        headers=["Goal", "Metric"],
        rows=[
            ["Adoption", "% of users who engage with Clara from the Brand Power Playbook entry point ('Ask Clara >')"],
            ["Engagement", "Reduction in out-of-scope queries over time (measured via graceful error response rate)"],
            ["Engagement", "% of Clara sessions that start from the Brand Power Playbook page vs other entry points"],
            ["Outcome", "User satisfaction with response relevance (qualitative / feedback signals)"],
            ["Outcome", "Zero misleading or irrelevant responses to non-Brand queries (verified via QA)"],
            ["Business Impact", "Users self-serve Brand insights without needing to re-query; reduced support requests for out-of-scope topics"],
        ]
    )

    _add_heading(doc, "Assumptions", level=2)
    _add_heading(doc, "Features in scope", level=3)
    _add_bullet_list(doc, [
        "Brand Power Playbook page as the primary Clara entry point",
        "UI labelling: Clara panel clearly identified as 'Clara Brand Guidance'",
        "Suggested prompts on landing are explicitly Brand-focused (e.g., brand performance, positioning, campaign insights)",
        "Graceful out-of-scope handling with the message: 'I'm sorry, I can only answer questions related to Brand Power Analytics.'",
        "Deep link redirect: users accessing a shared Clara conversation link who are not on the Brand Power Playbook page are redirected there, with the conversation loaded",
    ])
    _add_heading(doc, "Technical approach", level=3)
    _add_bullet_list(doc, [
        "UI changes to the IC component entry point and panel header to reflect Brand Guidance identity",
        "Prompt engineering / system prompt configuration to constrain Clara to Brand-related queries and generate the graceful error for out-of-scope inputs",
        "Deep link routing logic: detect current page context; if not Brand Power Playbook, redirect and then open the linked conversation",
    ])
    _add_heading(doc, "Data scope", level=3)
    _add_table(doc,
        headers=["Dimension", "Value"],
        rows=[
            ["Brands", "Brand Power Playbook brands"],
            ["Regions", "All regions where Brand Power Playbook is available"],
            ["Channels", "Web (OneWay application)"],
            ["Periods", "From release date onwards"],
            ["Granularity", "Per session / per query"],
        ]
    )

    _add_heading(doc, "Milestones", level=2)
    _add_table(doc,
        headers=["Milestone", "Date"],
        rows=[
            ["ADO Feature created", "—"],
            ["Development start", "TBD"],
            ["QA / validation", "TBD"],
            ["Release", "TBD"],
        ]
    )

    _add_heading(doc, "Requirements", level=2)
    _add_heading(doc, "Scenario 1 — Clara communicates Brand-specific purpose in UI", level=3)
    _add_bullet_list(doc, [
        "The 'Ask Clara >' entry point must be visible on the Brand Power Playbook page when the page loads.",
        "The Clara panel must clearly indicate its Brand Guidance purpose (e.g., labelled 'Clara Brand Guidance').",
        "Suggested prompts on the IC landing page must be explicitly Brand-focused (e.g., brand performance, positioning, campaign insights).",
    ])
    _add_heading(doc, "Scenario 2 — Clara handles out-of-scope queries intelligently", level=3)
    _add_bullet_list(doc, [
        "When a user submits a query unrelated to Brand insights (e.g., generic analytics, unrelated domains), Clara must recognise it as out-of-scope.",
        "Clara must respond with: 'I'm sorry, I can only answer questions related to Brand Power Analytics.'",
        "Clara must not generate misleading or irrelevant answers for out-of-scope queries.",
    ])
    _add_heading(doc, "Scenario 3 — Deep link enforces Brand context", level=3)
    _add_bullet_list(doc, [
        "When a user accesses a shared Clara conversation link and is not currently on the Brand Power Playbook page, the system must redirect them to the Brand Power Playbook page.",
        "After redirect, Clara must open with the shared conversation loaded.",
    ])

    _add_heading(doc, "Out of Scope", level=2)
    _add_heading(doc, "Short-term", level=3)
    _add_bullet_list(doc, [
        "Domain differentiation for other Clara deployments beyond the Brand Power Playbook / OneWay context",
        "Automatic reclassification of borderline queries (Clara either answers or gracefully declines — no partial answers)",
        "Analytics tracking of out-of-scope query volume (covered separately by Feature 156057)",
    ])
    _add_heading(doc, "Long-term", level=3)
    _add_bullet_list(doc, [
        "Multi-domain Clara configurations (e.g., a separate Clara instance scoped to a different analytics domain)",
    ])

    _add_heading(doc, "Design", level=2)
    _add_table(doc,
        headers=["Artefact", "Link"],
        rows=[
            ["Design file", "TBD"],
            ["IC panel branding mockup", "TBD"],
            ["Suggested prompts copy", "TBD"],
        ]
    )

    _add_heading(doc, "Open Questions", level=2)
    _add_table(doc,
        headers=["Question", "Answer", "Date Answered"],
        rows=[
            ["What is the exact label text for the Clara panel header? (e.g., 'Clara Brand Guidance' — confirm copy)", "—", "—"],
            ["What are the exact suggested prompts to display on the IC landing page?", "—", "—"],
            ["Who owns the out-of-scope detection logic — prompt engineering team or FE?", "—", "—"],
            ["Should the graceful error message be configurable or hardcoded?", "—", "—"],
            ["What happens if the deep link destination page is unavailable (e.g., user lacks access to Brand Power Playbook)?", "—", "—"],
            ["Does the deep link redirect preserve query parameters and conversation ID in all browsers / environments?", "—", "—"],
        ]
    )

    _add_heading(doc, "Change & Request Log", level=2)
    _add_table(doc,
        headers=["Request / Update", "Requestor / Personnel", "Date"],
        rows=[["Initial draft created from ADO 156063", "—", "2026-04-30"]]
    )

    _add_heading(doc, "Reference Links", level=2)
    _add_bullet_list(doc, [
        "ADO Feature 156063: https://dev.azure.com/ab-inbev-analytics",
        "Related: ADO 156057 (Product Analytics — tracks IC component interactions)",
        "Related: ADO 156061 (Share — conversation sharing)",
    ])

    _add_horizontal_rule(doc)

    _add_heading(doc, "Section 3 — Discovery → Delivery Checklist")

    _add_checklist_section(doc, "Phase 1 — Discovery Framing & Alignment", [
        "Problem statement defined: Clara lacks clear Brand Guidance identity, users submit out-of-scope queries",
        "Business intent confirmed: scoped, trusted Brand assistant with graceful boundaries",
        "Scope agreed: UI labelling, suggested prompts, out-of-scope handler, deep link redirect",
        "Ownership assigned (PM, UX, engineering, prompt engineering)",
        "ADO feature linked and prioritised",
    ], "Signed-off problem statement + scope boundary")

    _add_checklist_section(doc, "Phase 2 — Business Rules & Data Contract", [
        "Exact UI copy confirmed for Clara panel label and entry point",
        "Suggested prompt list agreed and signed off",
        "Out-of-scope response string confirmed: 'I'm sorry, I can only answer questions related to Brand Power Analytics.'",
        "Out-of-scope detection mechanism agreed (prompt engineering approach, keyword list, or classifier)",
        "Deep link redirect flow documented: trigger condition, redirect target, conversation load behaviour",
        "Edge case: user lacks Brand Power Playbook access — behaviour defined",
    ], "UI copy doc + out-of-scope handling spec")

    _add_checklist_section(doc, "Phase 3 — Discovery Validation (Users)", [
        "Users confirm they understand Clara's scope from the updated UI labels and prompts",
        "Out-of-scope query examples tested and graceful error confirmed as satisfactory",
        "Deep link redirect tested end-to-end with representative users",
    ], "User validation notes")

    _add_checklist_section(doc, "Phase 4 — HLR (High-Level Requirements)", [
        "MVP scope confirmed: 3 scenarios (UI branding, out-of-scope handler, deep link redirect)",
        "Non-goals documented (no multi-domain, no borderline query reclassification)",
        "Dependencies identified: IC component FE, prompt engineering / system prompt, routing / deep link logic",
        "Graceful error text confirmed as non-configurable (or flagged if config is needed)",
    ], "This document (HLR) approved")

    _add_checklist_section(doc, "Phase 5 — Design Handoff & Prototype Validation", [
        "IC panel header label design reviewed",
        "Suggested prompts layout reviewed against Brand Power Playbook page design",
        "Out-of-scope response UI treatment confirmed (same as standard Clara response vs. distinct styling)",
        "Deep link redirect UX reviewed (loading state, timing)",
    ], "Design sign-off from PO / UX lead")

    _add_checklist_section(doc, "Phase 6 — Delivery Prep (User Stories & Walkthrough)", [
        "User stories created in ADO for: UI label update, suggested prompts, out-of-scope handler, deep link redirect",
        "Acceptance criteria include: exact copy strings, out-of-scope trigger conditions, redirect flow",
        "Error states documented: what happens if redirect fails, if conversation ID is invalid",
        "QA test cases created for all 3 scenarios",
    ], "ADO stories with full AC, ready for sprint")

    _add_checklist_section(doc, "Phase 7 — Release & Communication", [
        "Release note drafted covering Clara brand differentiation changes",
        "Known limitations documented (e.g., out-of-scope detection may have edge cases on borderline queries)",
        "Stakeholders notified (PM, Brand analytics users, comms team)",
    ], "Release note + stakeholder comms sent")

    _add_checklist_section(doc, "Phase 8 — Post-Release Validation & Learning", [
        "Confirm 'Clara Brand Guidance' label and suggested prompts are live on Brand Power Playbook page",
        "Confirm out-of-scope queries return the correct graceful error message in production",
        "Confirm deep link redirect works correctly across environments and user access levels",
        "Monitor out-of-scope query rate and feedback signals",
        "v2 improvements identified (e.g., smarter out-of-scope detection, configurable prompts)",
    ], "Post-release validation checklist signed off")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

DOCS = [
    ("156057-product-analytics-owr.docx", _build_156057),
    ("156061-share-conversation-owr.docx", _build_156061),
    ("156063-domain-differentiation-owr.docx", _build_156063),
]


@app.command()
def main(
    out_dir: Path = typer.Option(DOCS_DIR, help="Output directory for generated .docx files"),
) -> None:
    """Generate OWR + HLR Word documents for ADO features 156057, 156061, 156063."""
    out_dir.mkdir(parents=True, exist_ok=True)

    for filename, builder in DOCS:
        out_path = out_dir / filename
        console.print(f"Building [bold]{filename}[/bold]...")
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        builder(doc)
        doc.save(str(out_path))
        console.print(f"  [green]✓ Saved → {out_path}[/green]")

    console.print(f"\n[bold green]Done.[/bold green] {len(DOCS)} documents written to {out_dir}/")


if __name__ == "__main__":
    app()
